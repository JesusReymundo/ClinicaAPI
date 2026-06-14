from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Margenes
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

PURPLE = RGBColor(0x6B, 0x21, 0xA8)
BLUE   = RGBColor(0x02, 0x84, 0xC7)
BLACK  = RGBColor(0, 0, 0)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x6B, 0x7B, 0x8D)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = PURPLE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = PURPLE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)
    return p

def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = GRAY
    r.font.size = Pt(10)
    return p

# ─── PORTADA ───────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INFORME TÉCNICO")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = PURPLE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Sistema de Gestión de Citas Médicas")
r2.bold = True
r2.font.size = Pt(18)
r2.font.color.rgb = BLACK

doc.add_paragraph()

info = [
    ("Institución",      "IDAT — Escuela de Tecnología"),
    ("Unidad didáctica", "Desarrollo de Servicios Web"),
    ("Ciclo",            "5to Ciclo"),
    ("Proyecto",         "ClinicaAPI — Evaluación Parcial"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(12)
    p.add_run(value).font.size = Pt(12)

doc.add_paragraph()

p_int = doc.add_paragraph()
p_int.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_int = p_int.add_run("Integrantes del Equipo")
r_int.bold = True
r_int.font.size = Pt(12)
r_int.font.color.rgb = PURPLE

integrantes = [
    "Jesus Reymundo Román",
    "Aldair Santos Cahuana",
    "Reyes Zarate Leomarc",
    "Ivan Zarate Soncco",
    "Crhistian Meza Cardenas",
    "Alexandee Morillo Campos",
]
for nombre in integrantes:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(nombre).font.size = Pt(11)

doc.add_page_break()

# ─── 1. DESCRIPCIÓN ────────────────────────────────────────
heading1("1. DESCRIPCIÓN DEL PROYECTO")

heading2("1.1 Contexto")
body("Una clínica privada necesita gestionar las citas médicas que los pacientes solicitan a diario. "
     "Anteriormente, las citas se anotaban manualmente en agendas físicas o se confirmaban por llamadas "
     "telefónicas sin un control unificado, generando confusión, pérdida de información y sobrecarga "
     "para el personal administrativo.")

heading2("1.2 Solución Implementada")
body("Se desarrolló una solución completa que consta de dos componentes principales:")
bullet("API RESTful con C# y .NET 8 conectada a una base de datos SQL Server (ClinicaDB) mediante "
       "Entity Framework Core 8, que expone todos los endpoints para la gestión clínica.")
bullet("Frontend web con React 18 + Vite que consume la API, con acceso diferenciado por rol "
       "(Administrador, Médico, Paciente) y una interfaz profesional tipo portal de clínica.")

body("La base de datos ClinicaDB contiene 23 tablas organizadas en 7 grupos funcionales, cubriendo "
     "desde la gestión de usuarios hasta facturación, historial clínico y auditoría.")

heading2("1.3 Tecnologías Utilizadas")

p_back = doc.add_paragraph()
r_back = p_back.add_run("Backend (API):")
r_back.bold = True
r_back.font.size = Pt(11)

bullet(".NET 8 / ASP.NET Core Web API")
bullet("Entity Framework Core 8 — ORM para SQL Server")
bullet("SQL Server (instancia: REYMUNDO\\SQL2025 · Base de datos: ClinicaDB)")
bullet("Swagger / OpenAPI — documentación interactiva en http://localhost:5024")
bullet("Inyección de dependencias + manejador global de excepciones")
bullet("JsonStringEnumConverter — permite enviar estados como texto (ej: 'Pendiente')")

doc.add_paragraph()
p_front = doc.add_paragraph()
r_front = p_front.add_run("Frontend (React):")
r_front.bold = True
r_front.font.size = Pt(11)

bullet("React 18 + Vite — SPA de alta performance")
bullet("React Router DOM v7 — enrutamiento por rol")
bullet("Axios — cliente HTTP para consumir la API")
bullet("CSS personalizado (~700 líneas) — diseño profesional tipo portal de clínica")
bullet("Control de acceso por roles (RBAC): Admin, Médico, Paciente via React Context")

doc.add_paragraph()
p_tools = doc.add_paragraph()
r_tools = p_tools.add_run("Herramientas:")
r_tools.bold = True
r_tools.font.size = Pt(11)

bullet("Postman — pruebas de endpoints")
bullet("GitHub — control de versiones")
bullet("Visual Studio Code / Visual Studio 2022")

# ─── 2. ESTRUCTURA ─────────────────────────────────────────
doc.add_page_break()
heading1("2. ESTRUCTURA DEL PROYECTO")

heading2("2.1 Arquitectura General")
body("El proyecto sigue una arquitectura de capas que separa claramente responsabilidades. "
     "La comunicación sigue el flujo: Frontend → API REST → Service → Repository (EF Core) → SQL Server.")

heading2("2.2 Estructura del Backend (ClinicaAPI)")

code_block("ClinicaAPI/")
code_block("├── Controllers/           → Exponen los endpoints REST (HTTP)")
code_block("│   ├── CitasController.cs")
code_block("│   ├── MedicosController.cs")
code_block("│   └── PacientesController.cs")
code_block("├── Services/              → Contienen la lógica de negocio")
code_block("│   ├── Interfaces/        → Contratos (ICitaService, etc.)")
code_block("│   ├── CitaService.cs")
code_block("│   ├── MedicoService.cs")
code_block("│   └── PacienteService.cs")
code_block("├── Models/                → Entidades del dominio (mapeadas a BD)")
code_block("│   ├── Cita.cs  |  Medico.cs  |  Paciente.cs")
code_block("│   ├── Usuario.cs  |  Rol.cs  |  Contacto.cs")
code_block("│   └── EstadoCita.cs  (enum: Pendiente/Confirmada/Cancelada/Completada/Anulada)")
code_block("├── DTOs/                  → Validaciones de entrada/salida")
code_block("│   ├── CitaDto.cs  |  MedicoDto.cs  |  PacienteDto.cs")
code_block("├── Data/")
code_block("│   └── ClinicaDbContext.cs → Contexto EF Core + DbSets")
code_block("├── Exceptions/            → Manejo centralizado de errores")
code_block("│   ├── AppExceptions.cs  |  GlobalExceptionHandler.cs")
code_block("└── Program.cs            → Configuración, DI, CORS, Swagger")

heading2("2.3 Estructura del Frontend (clinica-frontend)")

code_block("clinica-frontend/")
code_block("├── src/")
code_block("│   ├── api/")
code_block("│   │   └── clinicaApi.js   → Todas las llamadas HTTP (axios)")
code_block("│   ├── context/")
code_block("│   │   └── AuthContext.jsx → Autenticación y rol activo")
code_block("│   ├── pages/")
code_block("│   │   ├── Login.jsx       → Página de inicio de sesión")
code_block("│   │   ├── Dashboard.jsx   → Panel principal (por rol)")
code_block("│   │   ├── Pacientes.jsx   → CRUD de pacientes + panel lateral")
code_block("│   │   ├── Medicos.jsx     → CRUD de médicos + tarjetas")
code_block("│   │   └── Citas.jsx       → Gestión de citas + tabs por estado")
code_block("│   ├── App.jsx             → Layout, sidebar, rutas protegidas")
code_block("│   └── App.css             → Estilos globales (~700 líneas)")
code_block("└── vite.config.js")

heading2("2.4 Descripción de Capas del Backend")
layers = [
    ("Controllers:", "Reciben peticiones HTTP, validan el modelo y delegan al servicio. Devuelven respuestas con códigos HTTP correctos (200, 201, 400, 404). Todas las respuestas siguen el formato {success, data}."),
    ("Services:", "Lógica de negocio completa. Implementan interfaces para inyección de dependencias. Acceden a la base de datos mediante el DbContext de Entity Framework Core."),
    ("Models / DbContext:", "Entidades mapeadas a las 23 tablas de ClinicaDB. El DbContext configura relaciones, claves foráneas y restricciones. El enum EstadoCita define los 5 estados posibles."),
    ("DTOs:", "Separan la representación interna de lo que se expone al cliente. Validaciones con [Required], [StringLength], [EmailAddress], [RegularExpression]. PacienteResponseDto incluye Telefono, Email (de tabla Contactos), GrupoSanguineo y Alergias."),
    ("Exceptions:", "GlobalExceptionHandler captura excepciones no controladas y devuelve respuestas estructuradas en JSON. Tipos: NotFoundException (404) y BusinessException (400)."),
    ("CORS:", "Política FrontendReact configurada para permitir peticiones desde http://localhost:5173, :5174 y :3000."),
]
for title, desc in layers:
    p = doc.add_paragraph()
    r = p.add_run(title + " ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

# ─── 3. BASE DE DATOS ──────────────────────────────────────
doc.add_page_break()
heading1("3. BASE DE DATOS — CLINICADB")

heading2("3.1 Resumen")
body("La base de datos ClinicaDB (SQL Server, instancia REYMUNDO\\SQL2025) contiene 23 tablas "
     "organizadas en 7 grupos funcionales con más de 20 relaciones de clave foránea.")

heading2("3.2 Grupos de Tablas")

grupos = [
    ("Grupo A — Usuarios y Perfiles",
     "ROLES, USUARIOS, CONTACTOS",
     "Gestión de identidad y acceso. ROLES define los permisos; USUARIOS centraliza a todas las "
     "personas del sistema; CONTACTOS permite múltiples teléfonos/emails por usuario."),
    ("Grupo B — Personal Médico",
     "ESPECIALIDADES, MEDICOS, HORARIOSMEDICO, CONSULTORIOS",
     "Perfil y disponibilidad del médico. ESPECIALIDADES es catálogo; MEDICOS extiende USUARIOS "
     "con datos profesionales; HORARIOSMEDICO registra franjas de atención semanal."),
    ("Grupo C — Pacientes y Seguros",
     "TIPOASEGURADO, EMPRESAS, PACIENTES, SEGUROS",
     "Perfil clínico del paciente. PACIENTES extiende USUARIOS con grupo sanguíneo y alergias; "
     "SEGUROS registra múltiples pólizas por paciente."),
    ("Grupo D — Citas (Núcleo del sistema)",
     "ESTADOSCITA, CITAS, CANCELACIONESCITA",
     "CITAS es la tabla central que conecta pacientes con médicos. ESTADOSCITA es catálogo de "
     "estados. CANCELACIONESCITA registra motivo y responsable de cada cancelación."),
    ("Grupo E — Atención Médica",
     "TRIAJE, HISTORIALCLINICO, RECETAS",
     "Registro clínico de la consulta. TRIAJE almacena signos vitales; HISTORIALCLINICO acumula "
     "diagnósticos del paciente; RECETAS contiene la prescripción médica."),
    ("Grupo F — Medicamentos y Facturación",
     "MEDICAMENTOS, DETALLERECETA, FACTURAS",
     "MEDICAMENTOS es catálogo con stock; DETALLERECETA resuelve la relación N:N entre recetas y "
     "medicamentos; FACTURAS genera el comprobante de pago por cita."),
    ("Grupo G — Pagos y Auditoría",
     "METODOSPAGO, PAGOS, AUDITORIALOG, TIPOCONSULTA",
     "PAGOS registra transacciones múltiples por factura; AUDITORIALOG registra quién hizo qué "
     "y cuándo (INSERT/UPDATE/DELETE) para trazabilidad y seguridad."),
]

table_g = doc.add_table(rows=1 + len(grupos), cols=3)
table_g.style = 'Table Grid'
h_g = table_g.rows[0].cells
for i, txt in enumerate(["Grupo", "Tablas", "Propósito"]):
    h_g[i].text = txt
    h_g[i].paragraphs[0].runs[0].bold = True
    h_g[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    h_g[i].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(h_g[i], "6B21A8")
for i, (grp, tablas, prop) in enumerate(grupos):
    row = table_g.rows[i + 1].cells
    row[0].text = grp
    row[1].text = tablas
    row[2].text = prop
    for c in row:
        c.paragraphs[0].runs[0].font.size = Pt(9)
    if i % 2 == 0:
        for c in row:
            set_cell_bg(c, "F3E8FF")

doc.add_paragraph()

heading2("3.3 Estado Anulada — Diferenciación de Cancelación")
body("En la versión actualizada del sistema se distinguen dos tipos de cierre de cita:")

estados_tabla = [
    ["Cancelada",  "3", "La clínica o el médico cancela la cita (el paciente no llegó, el médico no pudo atender, etc.). Acción ejecutada por Admin o Médico. Endpoint: PATCH /api/citas/{id}/cancelar"],
    ["Anulada",    "5", "El propio paciente decide no asistir y anula la cita desde el portal. Acción exclusiva del rol Paciente. Endpoint: PATCH /api/citas/{id}/anular"],
    ["Completada", "4", "La consulta se realizó exitosamente. Solo Admin/Médico pueden marcar como completada."],
]

table_e = doc.add_table(rows=1 + len(estados_tabla), cols=3)
table_e.style = 'Table Grid'
h_e = table_e.rows[0].cells
for i, txt in enumerate(["Estado", "Código BD", "Significado y Responsable"]):
    h_e[i].text = txt
    h_e[i].paragraphs[0].runs[0].bold = True
    h_e[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    h_e[i].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(h_e[i], "6B21A8")
for i, row_data in enumerate(estados_tabla):
    row = table_e.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val
        row[j].paragraphs[0].runs[0].font.size = Pt(10)
    if i % 2 == 0:
        for c in row:
            set_cell_bg(c, "FFF7ED")

doc.add_paragraph()

# ─── 4. ENDPOINTS ──────────────────────────────────────────
doc.add_page_break()
heading1("4. ENDPOINTS DE LA API REST")

body("Todos los endpoints devuelven respuestas en formato JSON con la estructura: "
     "{ \"success\": true/false, \"data\": ... } o { \"success\": false, \"message\": \"...\" }")

def add_endpoint_table(title, rows):
    heading2(title)
    headers = ["Método", "Ruta", "Descripción", "Código HTTP"]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(hdr[i], "6B21A8")
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = cell_text
            row[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
        if r_idx % 2 == 0:
            for c in row:
                set_cell_bg(c, "F3E8FF")
    doc.add_paragraph()

add_endpoint_table("4.1 Pacientes (/api/Pacientes)", [
    ["GET",    "/api/Pacientes",          "Listar todos los pacientes (con teléfono, email, grupo sanguíneo)", "200 OK"],
    ["GET",    "/api/Pacientes/{id}",     "Buscar paciente por ID",                                           "200 / 404"],
    ["POST",   "/api/Pacientes",          "Registrar nuevo paciente",                                         "201 Created"],
    ["PUT",    "/api/Pacientes/{id}",     "Actualizar paciente",                                              "200 / 404"],
    ["DELETE", "/api/Pacientes/{id}",     "Eliminar paciente",                                                "200 / 404"],
])

add_endpoint_table("4.2 Médicos (/api/Medicos)", [
    ["GET",    "/api/Medicos",            "Listar todos los médicos con especialidad",                        "200 OK"],
    ["GET",    "/api/Medicos/{id}",       "Buscar médico por ID",                                            "200 / 404"],
    ["POST",   "/api/Medicos",            "Registrar nuevo médico",                                           "201 Created"],
    ["PUT",    "/api/Medicos/{id}",       "Actualizar médico",                                               "200 / 404"],
    ["DELETE", "/api/Medicos/{id}",       "Eliminar médico",                                                 "200 / 404"],
])

add_endpoint_table("4.3 Citas (/api/Citas)", [
    ["GET",    "/api/Citas",                      "Listar todas las citas",                                  "200 OK"],
    ["GET",    "/api/Citas/{id}",                 "Buscar cita por ID",                                      "200 / 404"],
    ["GET",    "/api/Citas/paciente/{id}",         "Citas de un paciente específico",                        "200 / 404"],
    ["GET",    "/api/Citas/medico/{id}",           "Citas de un médico específico",                          "200 / 404"],
    ["GET",    "/api/Citas/estado/{estado}",       "Filtrar por estado (Pendiente/Confirmada/Cancelada/Completada/Anulada)", "200 / 400"],
    ["POST",   "/api/Citas",                      "Registrar nueva cita (valida conflicto de horario ±30min)", "201 Created"],
    ["PUT",    "/api/Citas/{id}",                 "Actualizar cita",                                         "200 / 404"],
    ["PATCH",  "/api/Citas/{id}/cancelar",         "Cancelar cita (clínica/médico — paciente no llegó)",     "200 / 400"],
    ["PATCH",  "/api/Citas/{id}/anular",           "Anular cita (acción del propio paciente)",               "200 / 400"],
    ["DELETE", "/api/Citas/{id}",                 "Eliminar cita permanentemente",                           "200 / 404"],
])

heading2("4.4 Lógica de Negocio en Citas")
reglas = [
    "La fecha de la cita debe ser futura (no se puede agendar en el pasado)",
    "El médico no puede tener dos citas con menos de 30 minutos de diferencia (excluye Canceladas y Anuladas del cálculo)",
    "Una cita Cancelada o Anulada no puede modificarse ni volver a cancelarse",
    "Una cita Completada no puede cancelarse ni anularse",
    "Solo Admin y Médico pueden Confirmar, Completar o Cancelar una cita",
    "Solo el Paciente puede Anular su propia cita (endpoint /anular)",
    "El frontend muestra el botón correcto según el rol activo",
]
for r in reglas:
    bullet(r)

# ─── 5. FRONTEND REACT ─────────────────────────────────────
doc.add_page_break()
heading1("5. FRONTEND — PORTAL DE CLÍNICA (React)")

heading2("5.1 Acceso por Rol")
body("El sistema implementa control de acceso basado en roles (RBAC). Según el rol del usuario "
     "autenticado, el menú lateral, las opciones y los botones cambian dinámicamente:")

roles = [
    ["Administrador (Admin)", "admin123",    "Dashboard completo, gestión de Pacientes (CRUD), Médicos (CRUD), Citas (todas) con acciones de Confirmar/Completar/Cancelar. Ve estadísticas globales."],
    ["Médico (Doctor)",       "medico123",   "Panel propio, ver sus citas asignadas, ver lista de pacientes. Puede Confirmar, Completar y Cancelar citas. No puede crear nuevas citas."],
    ["Paciente",              "cita123",     "Ver sus propias citas, ver lista de médicos disponibles. Solo puede Anular sus citas (no Cancelar). No ve datos de otros pacientes."],
]

table_r = doc.add_table(rows=1 + len(roles), cols=3)
table_r.style = 'Table Grid'
h_r = table_r.rows[0].cells
for i, txt in enumerate(["Rol", "Credencial Demo", "Permisos"]):
    h_r[i].text = txt
    h_r[i].paragraphs[0].runs[0].bold = True
    h_r[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    h_r[i].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(h_r[i], "6B21A8")
for i, row_data in enumerate(roles):
    row = table_r.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val
        row[j].paragraphs[0].runs[0].font.size = Pt(10)
    if i % 2 == 0:
        for c in row:
            set_cell_bg(c, "F3E8FF")

doc.add_paragraph()

heading2("5.2 Páginas Principales")
paginas = [
    ("Login",      "Dos paneles: branding izquierdo + formulario derecho. Tarjetas de acceso rápido con credenciales demo por rol. Estadísticas del sistema (23 tablas, 3 roles, REST .NET 8)."),
    ("Dashboard",  "Tres vistas por rol: Admin (4 stat cards, tabla de citas recientes, acciones rápidas), Doctor (sus citas de hoy y próximas), Paciente (próxima cita destacada, historial)."),
    ("Pacientes",  "Tabla con avatar de iniciales, edad calculada, badge de grupo sanguíneo. Búsqueda por nombre/DNI/email. Filtro por grupo sanguíneo. Panel lateral con perfil completo + historial de citas. Formulario con sección de datos médicos (alergias, grupo sanguíneo)."),
    ("Médicos",    "Vista doble: tarjetas de especialidad (colores por área) + tabla. Filtro por especialidad. Panel lateral con perfil médico y todas sus citas. CRUD completo."),
    ("Citas",      "5 tabs: Todas / Pendientes / Confirmadas / Canceladas / Anuladas (con contadores). Acciones inline según estado y rol. Botones: Confirmar (Admin/Doctor), Completar, Cancelar (clínica), Anular (solo Paciente). Panel lateral con detalle de cita y acciones contextuales."),
]
for nombre, desc in paginas:
    p = doc.add_paragraph()
    r = p.add_run(nombre + ": ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = PURPLE
    p.add_run(desc).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

heading2("5.3 Comunicación con la API")
body("El frontend usa Axios con baseURL http://localhost:5024/api. Todas las respuestas de la API "
     "vienen envueltas en {success, data}, por lo que el cliente extrae la data con r.data.data. "
     "La función auxiliar d = r => r.data.data centraliza esta extracción en clinicaApi.js.")

code_block("// clinicaApi.js — funciones principales")
code_block("export const getPacientes   = () => api.get('/pacientes').then(d);")
code_block("export const getCitas       = () => api.get('/citas').then(d);")
code_block("export const cancelarCita   = id => api.patch(`/citas/${id}/cancelar`).then(d);")
code_block("export const anularCita     = id => api.patch(`/citas/${id}/anular`).then(d);")

# ─── 6. PRUEBAS POSTMAN ────────────────────────────────────
doc.add_page_break()
heading1("6. PRUEBAS FUNCIONALES CON POSTMAN")

body("Se realizaron pruebas de todos los endpoints de la API con Postman, verificando el "
     "correcto funcionamiento, validaciones y manejo de errores.")

note("Insertar capturas de pantalla de las pruebas realizadas en Postman.")

pruebas = [
    ("GET",    "/api/Pacientes",              "Listar pacientes con contactos",         "200 OK — array con teléfono y email"),
    ("POST",   "/api/Pacientes",              "Crear paciente",                         "201 Created"),
    ("PUT",    "/api/Pacientes/{id}",         "Actualizar grupo sanguíneo y alergias",  "200 OK"),
    ("DELETE", "/api/Pacientes/{id}",         "Eliminar paciente",                      "200 OK"),
    ("GET",    "/api/Medicos",                "Listar médicos con especialidad",         "200 OK"),
    ("POST",   "/api/Medicos",                "Crear médico",                           "201 Created"),
    ("GET",    "/api/Citas",                  "Listar todas las citas",                 "200 OK"),
    ("POST",   "/api/Citas",                  "Crear cita (fecha futura)",              "201 Created"),
    ("POST",   "/api/Citas",                  "Cita con conflicto de horario",          "400 — mensaje de error"),
    ("GET",    "/api/Citas/paciente/{id}",    "Citas de un paciente",                  "200 OK — filtrado correcto"),
    ("GET",    "/api/Citas/medico/{id}",      "Citas de un médico",                    "200 OK — filtrado correcto"),
    ("GET",    "/api/Citas/estado/Pendiente", "Filtrar por estado Pendiente",           "200 OK"),
    ("GET",    "/api/Citas/estado/Anulada",   "Filtrar por estado Anulada",            "200 OK"),
    ("PUT",    "/api/Citas/{id}",             "Actualizar cita",                        "200 OK"),
    ("PATCH",  "/api/Citas/{id}/cancelar",    "Cancelar cita (clínica)",               "200 OK — estado Cancelada"),
    ("PATCH",  "/api/Citas/{id}/anular",      "Anular cita (paciente)",                "200 OK — estado Anulada"),
    ("PATCH",  "/api/Citas/{id}/cancelar",    "Cancelar cita ya anulada (error)",      "400 — 'ya fue anulada por el paciente'"),
    ("DELETE", "/api/Citas/{id}",             "Eliminar cita",                          "200 OK"),
]

table2 = doc.add_table(rows=1 + len(pruebas), cols=4)
table2.style = 'Table Grid'
h2 = table2.rows[0].cells
for i, txt in enumerate(["Método", "Endpoint", "Descripción de la Prueba", "Resultado Esperado"]):
    h2[i].text = txt
    h2[i].paragraphs[0].runs[0].bold = True
    h2[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    h2[i].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(h2[i], "6B21A8")
for i, (met, ep, desc, res) in enumerate(pruebas):
    row = table2.rows[i + 1].cells
    row[0].text = met
    row[1].text = ep
    row[2].text = desc
    row[3].text = "✅ " + res
    for c in row:
        c.paragraphs[0].runs[0].font.size = Pt(9)
    if i % 2 == 0:
        for c in row:
            set_cell_bg(c, "F3E8FF")

doc.add_paragraph()

# ─── 7. REPOSITORIO ────────────────────────────────────────
heading1("7. REPOSITORIO EN GITHUB")
body("El código fuente del proyecto (backend y frontend) está disponible en el siguiente repositorio público:")

p_repo = doc.add_paragraph()
r_repo = p_repo.add_run("https://github.com/JesusReymundo/ClinicaAPI")
r_repo.bold = True
r_repo.font.size = Pt(12)
r_repo.font.color.rgb = PURPLE

heading2("Instrucciones de instalación y ejecución")

p_b = doc.add_paragraph()
r_b = p_b.add_run("Backend (API):")
r_b.bold = True; r_b.font.size = Pt(11)

code_block("git clone https://github.com/JesusReymundo/ClinicaAPI.git")
code_block("cd ClinicaAPI")
code_block("dotnet restore")
code_block("dotnet run")
code_block("# → API disponible en http://localhost:5024")
code_block("# → Swagger UI en http://localhost:5024 (raíz)")

doc.add_paragraph()

p_f = doc.add_paragraph()
r_f = p_f.add_run("Frontend (React):")
r_f.bold = True; r_f.font.size = Pt(11)

code_block("cd clinica-frontend")
code_block("npm install")
code_block("npm run dev")
code_block("# → Portal web en http://localhost:5173 (o :5174 si ya está ocupado)")

doc.add_paragraph()
body("Credenciales de acceso demo:")
bullet("admin / admin123  →  Rol Administrador (acceso completo)")
bullet("doctor / medico123  →  Rol Médico (sus citas y pacientes)")
bullet("paciente / cita123  →  Rol Paciente (solo sus citas, puede anular)")

# ─── 8. ROLES DEL EQUIPO ───────────────────────────────────
doc.add_page_break()
heading1("8. ROLES DEL EQUIPO")

team = [
    ["Jesus Reymundo Román",    "Líder técnico / Full Stack",  "Arquitectura general, Módulo Citas (incluyendo estado Anulada), SQL Server + EF Core, CORS, Frontend React completo, Diagrama ER"],
    ["Aldair Santos Cahuana",   "Backend",                     "Módulo Pacientes, DTOs con GrupoSanguíneo y Alergias, fix de Contactos en PacienteService"],
    ["Reyes Zarate Leomarc",    "Backend",                     "Módulo Médicos, modelos de dominio, especialidades"],
    ["Ivan Zarate Soncco",      "Backend",                     "Validaciones, DTOs de entrada y salida, testing"],
    ["Crhistian Meza Cardenas", "Backend",                     "Manejo centralizado de excepciones, GlobalExceptionHandler"],
    ["Alexandee Morillo Campos","Documentación / QA",          "Swagger, pruebas Postman, README, documentación técnica"],
]

table3 = doc.add_table(rows=1 + len(team), cols=3)
table3.style = 'Table Grid'
h3 = table3.rows[0].cells
for i, txt in enumerate(["Integrante", "Rol", "Responsabilidad"]):
    h3[i].text = txt
    h3[i].paragraphs[0].runs[0].bold = True
    h3[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    h3[i].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(h3[i], "6B21A8")
for i, row_data in enumerate(team):
    row = table3.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val
        row[j].paragraphs[0].runs[0].font.size = Pt(10)
    if i % 2 == 0:
        for c in row:
            set_cell_bg(c, "F3E8FF")

doc.add_paragraph()

# ─── 9. CONCLUSIONES ───────────────────────────────────────
heading1("9. CONCLUSIONES")

conclusiones = [
    "Se implementó una API RESTful completa con .NET 8 y Entity Framework Core 8 conectada a una base de datos SQL Server real con 23 tablas relacionadas, superando la fase inicial de almacenamiento en memoria.",
    "La distinción entre el estado Cancelada (acción de la clínica) y Anulada (acción del paciente) mejora la trazabilidad y permite análisis de indicadores de calidad más precisos.",
    "El frontend React con control de acceso por roles demuestra cómo consumir una API RESTful de manera segura y eficiente, con experiencia de usuario diferenciada según el perfil del actor.",
    "La arquitectura de capas (Controllers → Services → EF Core → SQL Server) facilita el mantenimiento, las pruebas unitarias y la escalabilidad futura del sistema.",
    "El uso de CORS, DTOs con validaciones, JsonStringEnumConverter y un manejador global de excepciones refleja buenas prácticas en el desarrollo de servicios web profesionales.",
]
for i, c in enumerate(conclusiones, 1):
    bullet(c, bold_prefix=f"{i}.")

doc.add_paragraph()

out = r"d:\IDAT\5TO CICLO\Desarrollo de Servicios Web\ClinicaAPI\Informe_Tecnico_ClinicaAPI.docx"
doc.save(out)
print(f"Documento generado: {out}")
